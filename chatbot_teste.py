import streamlit as st 
import openai
import os
from streamlit_chat import message as msg
import docx # criar e editar doc do word
import io # auxiliar criar, editar e salvar arquivo do word dentro do streamlit
from datetime import datetime



openai.api_key = os.getenv("SENHA_OPEN_AI")


st.title("Chat teste com ChatGPT Turbo")
st.write("***")

# trecho de código a seguir - session_state - guardar o histórico do chatbot
if 'hst_conversa' not in st.session_state:
    st.session_state.hst_conversa = []


pergunta = st.text_input("Digite sua pergunta: ")
btn_enviar = st.button("Enviar pergunta")

if btn_enviar:
    st.session_state.hst_conversa.append({"role": "user", "content": pergunta})
    # envio da solicitação da API da OpenAi
    retorno_openai = openai.ChatCompletion.create(
        model = "gpt-3.5-turbo",
        messages = st.session_state.hst_conversa, # list com cada elemento sendo um dic
        max_token = 500,
        n=1
    )
    # trecho de código a seguir retorna um dic com o papel e o conteúdo e armazena a resposta da openai
    st.session_state.hst_conversa.append(
        {"role": "assistant",
         "content": retorno_openai['choices'][0]['content']}
        )

# exibir somente quando a lista de elementos de memória for maior que 0
if len(st.session_state.hst_conversa) > 0:
    for i in range(len(st.session_state.hst_conversa)): 
        # exibição de conteúdo - posição par =  usário, posisão impar = resposta da ia
        if i % 2 == 0:
            msg("Você: " + st.session_state.hst_conversa[i]['content'], is_user=True)
        else:
            msg("Resposta IA: " + st.session_state.hst_conversa[i]['content'])
 
# salvar o conteúdo da conversa em um doc word
if len(st.session_state.hst_conversa) > 0:
    btn_salvar = st.button("Salvar o contúdo")
    if btn_salvar:
        trabalho = io.BytesIO() # aloca a variável com espaço em bytes para salvar o documento
        documento = docx.Document() # gerar documento do word
        # Adiciona o heading com data e hora
        now = datetime.now()
        formatted_date = now.strftime("%Y-%m-%d %H:%M:%S")
        documento.add_heading(f"Conteúdo Gerado - {formatted_date}", level=1)
        
        # criação do botão para salvar o conteúdo
        for i in range(len(st.session_state.hst_conversa)):
            if i % 2 == 0:
                documento.add_heading("Pergunta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])
            else:
                documento.add_heading("Pergunta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])
                
        documento.save(trabalho)
        st.download_button(label="Clique aqui para salvar o conteúdo da conversa",
                           data=trabalho,
                           file_name="",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
   
    